from docx import Document
from docx.shared import Pt
from docx.oxml.ns import nsdecls
from docx.oxml import OxmlElement
import re
from docx2pdf import convert
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.enum.section import WD_SECTION_START
from docx.shared import Cm, Inches
from docx.shared import RGBColor
from docx.oxml.ns import qn

# Function definitions


def start_each_heading_from_new_page(doc):
    first_heading_found = False
    i = 0
    while i < len(doc.paragraphs):
        paragraph = doc.paragraphs[i]
        is_heading(paragraph, doc)
        print(i)
        i += 1


def clear_paragraph(paragraph):
    # Clear paragraph content
    paragraph.clear()
    paragraph.text = ""
    for run in paragraph.runs:
        run.clear()

def remove_tables_associated_with_paragraph(doc, paragraph_index):
    # Remove tables associated with the paragraph at given index
    tables_to_remove = []
    for table in doc.tables:
        if table._element.getparent().index(table._element) > paragraph_index:
            break
        tables_to_remove.append(table)
    for table in tables_to_remove:
        table._element.getparent().remove(table._element)


def remove_heading_and_content(doc, headings_to_remove):
    for heading_to_remove in headings_to_remove:
        heading_found = False
        
        for i, paragraph in enumerate(doc.paragraphs):
            if is_heading(paragraph):
                heading_text = paragraph.text.strip().split('\n')[0]
                if heading_text == heading_to_remove:
                    heading_found = True
                    print(f"Found heading '{heading_text}' at index {i}")
                    
                    # Remove content under the heading
                    remove_next = True
                    for j in range(i + 1, len(doc.paragraphs)):
                        if is_heading(doc.paragraphs[j]):
                            break
                        if remove_next:
                            clear_paragraph(doc.paragraphs[j])
                            # remove_tables_associated_with_paragraph(doc, j)
                    
                    # Clear heading text
                    clear_paragraph(doc.paragraphs[i])
                    
                    # Remove tables associated with the heading paragraph
                    remove_tables_associated_with_paragraph(doc, i)
                    
                    # Remove the heading paragraph itself
                    p_element = paragraph._element
                    p_element.getparent().remove(p_element)
                    
                    # Stop processing after this heading
                    break
        
        if heading_found:
            print(f"Removed heading '{heading_to_remove}' and its content.")
            replace_heading_numbering(doc)
        else:
            print(f"Heading '{heading_to_remove}' not found or not recognized as a heading.")



def add_page_border(doc, border_space=20):
    for section in doc.sections:
        section_start = section._sectPr
        border_element = OxmlElement('w:pgBorders')
        border_element.set(qn('w:offsetFrom'), 'page')

        for border_type in ['top', 'bottom', 'left', 'right']:
            border = OxmlElement(f'w:{border_type}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '10')
            border.set(qn('w:space'), str(border_space))  # Set the space between border and page edge
            border.set(qn('w:color'), 'auto')
            border_element.append(border)

        section_start.append(border_element)

def set_page_size_to_a4(doc):
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)

    # Ensure all sections have the same page size
    for section in doc.sections:
        section.start_type = WD_SECTION_START.NEW_PAGE



def add_header_with_image_size(doc, image_path, width_cm, height_cm):
    section = doc.sections[0]  # Get the first section of the document

    # Add header
    header = section.header
    header_paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()

    # Add image to header with specific width and height
    run = header_paragraph.add_run()
    picture = run.add_picture(image_path)
    picture.width = Cm(width_cm)
    picture.height = Cm(height_cm)

    # Align header content
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT



def add_footer_with_page_number(doc):
    section = doc.sections[0]  # Get the first section of the document

    # Create footer
    footer = section.footer
    footer_paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()

    # Add page number field to footer
    run = footer_paragraph.add_run()
    fld_char = OxmlElement('w:fldChar')
    fld_char.set(qn('w:fldCharType'), 'begin')
    run._r.append(fld_char)

    run = footer_paragraph.add_run()
    fld_simple = OxmlElement('w:instrText')
    fld_simple.text = 'PAGE'
    run._r.append(fld_simple)

    run = footer_paragraph.add_run()
    fld_char = OxmlElement('w:fldChar')
    fld_char.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_char)

    # Align page number to the right
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT





def is_heading(paragraph):
    # Check if any line in the paragraph starts with an integer followed by a dot and a space or an integer followed by a space
    for line in paragraph.text.splitlines():
        if re.match(r'^\d+(\.| ) ', line):
            # Check if any run in the paragraph is bold and size is >= 12
            for run in paragraph.runs:
                if line in run.text and run.bold and run.font.size and run.font.size.pt >= 12:
                    return True
    return False

def remove_content_above_first_heading(doc):
    first_heading_index = None

    # Find the first heading
    for i, paragraph in enumerate(doc.paragraphs):
        if is_heading(paragraph):
            first_heading_index = i
            break

    if first_heading_index is not None:
        # Remove all elements before the first heading
        elements_to_remove = []
        for i, element in enumerate(doc.element.body):
            if i < first_heading_index:
                elements_to_remove.append(element)
            else:
                break
        
        for element in elements_to_remove:
            element.getparent().remove(element)



def prompt_for_headings_to_remove(headings):
    print("Available Headings:")
    for idx, heading in enumerate(headings, start=1):
        print(f"{idx}. {heading}")

    choices = input("Enter the numbers of the headings you want to remove separated by commas (e.g., '1, 3') or enter '0' to skip: ")
    if choices.strip() == '0':
        return []

    selected_indices = []
    try:
        selected_indices = [int(idx.strip()) for idx in choices.split(',')]
        invalid_indices = [idx for idx in selected_indices if idx < 1 or idx > len(headings)]
        if invalid_indices:
            print(f"Invalid choices: {', '.join(map(str, invalid_indices))}. Please enter valid numbers.")
            return prompt_for_headings_to_remove(headings)
        return [headings[idx - 1] for idx in selected_indices]
    except ValueError:
        print("Invalid input. Please enter numbers separated by commas.")
        return prompt_for_headings_to_remove(headings)
    


def Format_doc(doc):
    paragraphs_to_remove = []

    # This Function Is Use for Start Each Heading in New Page

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
        # Collect runs to modify in a separate list to avoid modifying while iterating
        runs_to_modify = []
        for run in paragraph.runs:
            for br in run._element.findall(qn('w:br')):
                if br.get(qn('w:type')) == 'page':
                    runs_to_modify.append(run)
                    break  # No need to check further if a page break is found
        
        # Merge content with adjacent paragraphs
        for run in runs_to_modify:
            run_index = paragraph.runs.index(run)
            
            if run_index < len(paragraph.runs) - 1:
                # Merge text with the next run in the current paragraph
                next_run = paragraph.runs[run_index + 1]
                next_run.text = run.text + next_run.text
                run.clear()
            else:
                # Merge text with the next paragraph
                next_paragraph = paragraph._element.getnext()
                if next_paragraph is not None:
                    next_paragraph.text = run.text + next_paragraph.text
                    run.clear()
                    paragraphs_to_remove.append(paragraph)
    
    # Remove empty paragraphs after merging
    for paragraph in paragraphs_to_remove:
        p_element = paragraph._element
        p_element.getparent().remove(p_element)

    # Remove excessive spaces
    for paragraph in doc.paragraphs:
        if len(paragraph.runs) > 0:
            for run in paragraph.runs:
                text = run.text
                if text.count(' ') >= 3:
                    # Replace more than 3 consecutive spaces with single space
                    modified_text = ' '.join(text.split())
                    run.text = modified_text
    
    # Remove entirely empty paragraphs
    empty_paragraphs = [p for p in doc.paragraphs if p.text.strip() == '']
    for paragraph in empty_paragraphs:
        p_element = paragraph._element
        p_element.getparent().remove(p_element)
    
    for table in doc.tables:
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER




def replace_heading_numbering(doc):
    main_number = 0  # Initialize main heading number
    sub_number = 0  # Initialize subheading number
    current_main_number = 0  # Track current main heading number
    current_sub_number = 0  # Track current subheading number
    
    for paragraph in doc.paragraphs:
        if is_heading(paragraph):
            # Determine heading level and adjust numbering accordingly
            if re.match(r'^\d+\. ', paragraph.text.strip()):
                current_main_number += 1
                main_number = current_main_number  # Update main_number to current main heading number
                sub_number = 0  # Reset sub-number for new main heading
            elif re.match(r'^\d+\.\d+', paragraph.text.strip()):  # Detecting subheading format like 1.1, 1.2
                current_sub_number += 1
                sub_number = current_sub_number  # Update sub_number to current subheading number
            
            # Construct new heading text with updated numbering
            if re.match(r'^\d+\. ', paragraph.text.strip()):
                new_heading_text = "{}. {}".format(main_number, re.sub(r'^\d+\. ', '', paragraph.text.strip()))
            elif re.match(r'^\d+\.\d+', paragraph.text.strip()):  # Detecting subheading format like 1.1, 1.2
                new_heading_text = "{}.{}. {}".format(main_number, sub_number, re.sub(r'^\d+\.\d+\s*', '', paragraph.text.strip()))
            else:
                new_heading_text = paragraph.text.strip()
            
            # Replace heading text with updated numbering
            paragraph.text = new_heading_text
            paragraph.style = 'Heading 1'
            
            # Apply bold and size 14 to heading
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(255, 0, 0)
        
        elif is_subheading(paragraph):
            # If it's a subheading (e.g., 1.1, 1.2, ...), update its numbering
            sub_number += 1  # Increment sub_number within the current main heading
            
            # Split the text to extract the heading number and the text after it
            heading_number, heading_text = re.split(r'\s+', paragraph.text.strip(), 1)
            
            # Construct new subheading text with updated numbering
            new_subheading_text = "{}.{}. {}".format(main_number, sub_number, heading_text.strip())
            paragraph.text = new_subheading_text
            paragraph.style = 'Heading 2'  
            
            # Apply bold and size 12 to subheading
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0, 0, 255)

def is_subheading(paragraph):
    # Check if paragraph starts with a subheading number pattern (e.g., 1.1, 2.3, ...)
    return re.match(r'^\d+\.\d+', paragraph.text.strip())


def replace_table_number_series(doc):
    current_table_number = 0  # Initialize current table number

    for table in doc.tables:
        current_table_number += 1
        # Replace or update table number in the document
        table_title = f"Table {current_table_number}. "
        new_table_title = f"Table {current_table_number}. "
        
        for paragraph in doc.paragraphs:
            if table_title in paragraph.text:
                paragraph.text = paragraph.text.replace(table_title, new_table_title)
                break


