from docx import Document
from docx.shared import Pt
from docx.oxml.ns import nsdecls
from docx.oxml import OxmlElement
import re
from docx2pdf import convert
from pdf2docx import Converter
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.shared import Cm, Inches
from docx.shared import RGBColor
from docx.oxml.ns import qn
import os
import shutil

# Function definitions
def remove_content_under_heading(doc, headings_to_remove):
    for heading_to_remove in headings_to_remove:
        remove_next = False
        for i, paragraph in enumerate(doc.paragraphs):
            if is_heading(paragraph) and paragraph.text.strip() == heading_to_remove:
                remove_next = True
                # Remove the heading itself
                p_element = paragraph._element
                p_element.getparent().remove(p_element)
                continue  # Skip removing the heading's content
            elif remove_next:
                if is_heading(paragraph):
                    # If another heading is encountered, stop removing content
                    break
                else:
                    # Remove tables associated with the paragraph
                    tables_to_remove = []
                    for j, table in enumerate(doc.tables):
                        if table._element.getparent().index(table._element) > i:
                            break
                        tables_to_remove.append(table)
                    for table in tables_to_remove:
                        table._element.getparent().remove(table._element)
                    
                    # Remove the paragraph itself
                    p_element = paragraph._element
                    p_element.getparent().remove(p_element)
        
        # After removing content under the heading, renumber headings
        replace_heading_numbering(doc)


def add_page_border(doc, border_space=20):
    for section in doc.sections:
        section_start = section._sectPr
        border_element = OxmlElement('w:pgBorders')
        border_element.set(qn('w:offsetFrom'), 'page')

        for border_type in ['top', 'bottom', 'left', 'right']:
            border = OxmlElement(f'w:{border_type}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '10')
            border.set(qn('w:space'), str(border_space))  # Set the space between border and page edge
            border.set(qn('w:color'), 'auto')
            border_element.append(border)

        section_start.append(border_element)

def set_page_size_to_a4(doc):
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)

    # Ensure all sections have the same page size
    for section in doc.sections:
        section.start_type = WD_SECTION_START.NEW_PAGE



def add_header_with_image_size(doc, image_path, width_cm, height_cm):
    section = doc.sections[0]  # Get the first section of the document

    # Add header
    header = section.header
    header_paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()

    # Add image to header with specific width and height
    run = header_paragraph.add_run()
    picture = run.add_picture(image_path)
    picture.width = Cm(width_cm)
    picture.height = Cm(height_cm)

    # Align header content
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT



def add_footer_with_page_number(doc):
    section = doc.sections[0]  # Get the first section of the document

    # Create footer
    footer = section.footer
    footer_paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()

    # Add page number field to footer
    run = footer_paragraph.add_run()
    fld_char = OxmlElement('w:fldChar')
    fld_char.set(qn('w:fldCharType'), 'begin')
    run._r.append(fld_char)

    run = footer_paragraph.add_run()
    fld_simple = OxmlElement('w:instrText')
    fld_simple.text = 'PAGE'
    run._r.append(fld_simple)

    run = footer_paragraph.add_run()
    fld_char = OxmlElement('w:fldChar')
    fld_char.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_char)

    # Align page number to the right
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT





def is_heading(paragraph):
    # Check if paragraph starts with an integer followed by a dot and a space or an integer followed by a space
    if not re.match(r'^\d+(\.| ) ', paragraph.text):
        return False

    # Check if text is bold and size is >= 12
    for run in paragraph.runs:
        if run.bold and run.font.size and run.font.size.pt >= 12:
            return True
    return False



def replace_heading_numbering(doc):
    main_number = 0  # Initialize main heading number
    sub_number = 0  # Initialize subheading number
    current_main_number = 0  # Track current main heading number
    current_sub_number = 0  # Track current subheading number
    
    for paragraph in doc.paragraphs:
        if is_heading(paragraph):
            # Determine heading level and adjust numbering accordingly
            if re.match(r'^\d+\. ', paragraph.text.strip()):
                current_main_number += 1
                main_number = current_main_number  # Update main_number to current main heading number
                sub_number = 0  # Reset sub-number for new main heading
            elif re.match(r'^\d+\.\d+', paragraph.text.strip()):  # Detecting subheading format like 1.1, 1.2
                current_sub_number += 1
                sub_number = current_sub_number  # Update sub_number to current subheading number
            
            # Construct new heading text with updated numbering
            if re.match(r'^\d+\. ', paragraph.text.strip()):
                new_heading_text = "{}. {}".format(main_number, re.sub(r'^\d+\. ', '', paragraph.text.strip()))
            elif re.match(r'^\d+\.\d+', paragraph.text.strip()):  # Detecting subheading format like 1.1, 1.2
                new_heading_text = "{}.{}. {}".format(main_number, sub_number, re.sub(r'^\d+\.\d+\s*', '', paragraph.text.strip()))
            else:
                new_heading_text = paragraph.text.strip()
            
            # Replace heading text with updated numbering
            paragraph.text = new_heading_text
            
            # Apply bold and size 14 to heading
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(14)
        
        elif is_subheading(paragraph):
            # If it's a subheading (e.g., 1.1, 1.2, ...), update its numbering
            sub_number += 1  # Increment sub_number within the current main heading
            
            # Split the text to extract the heading number and the text after it
            heading_number, heading_text = re.split(r'\s+', paragraph.text.strip(), 1)
            
            # Construct new subheading text with updated numbering
            new_subheading_text = "{}.{}. {}".format(main_number, sub_number, heading_text.strip())
            paragraph.text = new_subheading_text
            
            # Apply bold and size 12 to subheading
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(12)

def is_subheading(paragraph):
    # Check if paragraph starts with a subheading number pattern (e.g., 1.1, 2.3, ...)
    return re.match(r'^\d+\.\d+', paragraph.text.strip())


def replace_table_number_series(doc):
    current_table_number = 0  # Initialize current table number

    for table in doc.tables:
        current_table_number += 1
        # Replace or update table number in the document
        table_title = f"Table {current_table_number}. "
        new_table_title = f"Table {current_table_number}. "
        
        for paragraph in doc.paragraphs:
            if table_title in paragraph.text:
                paragraph.text = paragraph.text.replace(table_title, new_table_title)
                break




def remove_content_above_first_heading(doc):
    first_heading_index = None

    # Find the first heading
    for i, paragraph in enumerate(doc.paragraphs):
        if is_heading(paragraph):
            first_heading_index = i
            break

    if first_heading_index is not None:
        # Remove all elements before the first heading
        elements_to_remove = []
        for i, element in enumerate(doc.element.body):
            if i < first_heading_index:
                elements_to_remove.append(element)
            else:
                break
        
        for element in elements_to_remove:
            element.getparent().remove(element)

def remove_heading(doc, heading_text):
    found_heading = False
    for i, paragraph in enumerate(doc.paragraphs):
        if is_heading(paragraph) and paragraph.text.strip() == heading_text:
            found_heading = True
            print(f"Found heading '{heading_text}' at index {i}")

            # Clear heading text
            doc.paragraphs[i].clear()
            doc.paragraphs[i].text = ""

            # Remove consecutive empty paragraphs
            for j in range(i + 1, len(doc.paragraphs)):
                if doc.paragraphs[j].text.strip() == "":
                    print(f"Removing empty paragraph at index {j}")
                    doc.paragraphs[j].clear()
                    doc.paragraphs[j].text = ""
                else:
                    break
            
            break
    
    if found_heading:
        replace_heading_numbering(doc)
    else:
        print(f"Heading '{heading_text}' not found or not recognized as a heading.")


def prompt_for_headings_to_remove(headings):
    print("Available Headings:")
    for idx, heading in enumerate(headings, start=1):
        print(f"{idx}. {heading}")

    choices = input("Enter the numbers of the headings you want to remove separated by commas (e.g., '1, 3') or enter '0' to skip: ")
    if choices.strip() == '0':
        return []

    selected_indices = []
    try:
        selected_indices = [int(idx.strip()) for idx in choices.split(',')]
        invalid_indices = [idx for idx in selected_indices if idx < 1 or idx > len(headings)]
        if invalid_indices:
            print(f"Invalid choices: {', '.join(map(str, invalid_indices))}. Please enter valid numbers.")
            return prompt_for_headings_to_remove(headings)
        return [headings[idx - 1] for idx in selected_indices]
    except ValueError:
        print("Invalid input. Please enter numbers separated by commas.")
        return prompt_for_headings_to_remove(headings)
    












def Get_New_PDF(Docx_path,Header_image_path,Folder_name):
    try:
        os.makedirs(Folder_name,exist_ok=True)
        width_cm = 5.00
        height_cm = 1.44 
        doc = Document(Docx_path)
        set_page_size_to_a4(doc)
        add_page_border(doc)
       

        # Remove content above the first heading
        remove_content_above_first_heading(doc)

        # Add header and footer
        add_header_with_image_size(doc, Header_image_path,width_cm,height_cm )

        # Replace heading numbering in series and apply bold and size 14
        replace_heading_numbering(doc)

        # Extract all headings
        headings = [paragraph.text.strip() for paragraph in doc.paragraphs if is_heading(paragraph)]

        # Prompt user for headings to remove
        headings_to_remove = prompt_for_headings_to_remove(headings)
        if headings_to_remove:
            remove_content_under_heading(doc=doc,headings_to_remove=headings_to_remove)

            add_footer_with_page_number(doc)

            replace_table_number_series(doc)

            # Save the modified document
            output_docx = Folder_name+'/modified_SDCIT.docx'
            output_pdf = 'Final.pdf'
            doc.save(output_docx)


            convert(output_docx, output_pdf)

            print(f"Headings removed and numbering adjusted in '{output_docx}'")
        else:
            print("No headings selected for removal. Exiting.")
        # shutil.rmtree(Folder_name)
    except Exception as e:
        print(e)



Get_New_PDF(Docx_path='SDCIT.docx',
            Header_image_path='flexon_logo.png',
            Folder_name='Temp_folder'
            )