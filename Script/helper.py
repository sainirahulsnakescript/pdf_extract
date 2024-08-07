from docx.shared import Pt, Cm,RGBColor
from docx.oxml import OxmlElement, ns
import re
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# Watermark requirement
from PyPDF2 import PdfReader, PdfWriter
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.colors import black

from io import BytesIO
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from pdf2docx import Converter
import fitz
from docx.oxml.ns import nsdecls
import docx
import win32com.client
import os


def convert_pdf_to_docx(file_path):
    try:
        cv = Converter(file_path)
        docx_path = file_path.replace(".pdf", ".docx")
        cv.convert(docx_path)
        cv.close()
        print(f"Conversion completed for {file_path}")
        return docx_path
    except Exception as e:
        print(f"Error converting {file_path}: {e}")

def is_heading(paragraph):
    for line in paragraph.text.splitlines():
        if re.match(r'^\d+(\.| |\t).*', line):
            if paragraph.style.name == 'Heading 1':
                return True
            # Check if any run in the paragraph is bold and size is >= 12
            # for run in paragraph.runs:
            #     if line in run.text and run.bold and run.font.size and run.font.size.pt >= 12:
            #         return True
    return False

def is_subheading(paragraph):
    for line in paragraph.text.splitlines():
        if re.match(r'^\d+\.\d+(\.| |\t)', line):
            if paragraph.style.name == 'Heading 2':
                    return True
            elif paragraph.style.name == 'Heading 3':
                return False
            for run in paragraph.runs:
                if run.text in line and run.bold:
                    return True

    # Check if paragraph starts with a subheading number pattern (e.g., 1.1, 2.3, ...)
    return False

def is_subheading_heading(paragraph):
    for line in paragraph.text.splitlines():
        if re.match(r'^\d+\.\d+', line):
            return True


def is_sub_subheading(paragraph):
    for line in paragraph.text.splitlines():
        if re.match(r'^\d+\.\d+\.\d+', line):
            if paragraph.style.name == 'Heading 3':
                return True
            for run in paragraph.runs:
                if run.text in line or run.bold:
                    return True
    return False





def remove_content_above_first_heading(doc):
    first_heading_index = None

    # Find the index of the first heading
    for i, paragraph in enumerate(doc.paragraphs):
        if is_heading(paragraph):
            first_heading_index = i
            break

    if first_heading_index is not None:
        # Remove paragraphs before the first heading
        for i in range(first_heading_index):
            doc.paragraphs[i].clear()  # Clear the content of the paragraph

        # Remove any additional content like tables, images, etc., if needed
        for element in doc.element.body[:first_heading_index]:
            element.getparent().remove(element)


def lock_table(table, lock=True):
    tbl = table._element
    tblPr = tbl.xpath("w:tblPr")[0]
    tblLocked = OxmlElement('w:tblLocked')
    tblLocked.set(ns.qn('w:val'), 'true' if lock else 'false')
    tblPr.append(tblLocked)


def Format_doc(doc):
    paragraphs_to_remove = []

    for section in doc.sections:
        margin = 1.27
        section.top_margin = Cm(margin)
        section.bottom_margin = Cm(margin)
        section.left_margin = Cm(margin)
        section.right_margin = Cm(margin)

        # Ensure all sections start on a new page
        section.start_type = WD_SECTION_START.NEW_PAGE

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            # Set the paragraph indentation
        paragraph.paragraph_format.left_indent = Cm(1.5)
        paragraph.paragraph_format.right_indent = Cm(1.5)
        # Collect runs to modify in a separate list to avoid modifying while iterating
        runs_to_modify = []
        for run in paragraph.runs:
            for br in run._element.findall(qn('w:br')):
                if br.get(qn('w:type')) == 'page':
                    runs_to_modify.append(run)
                    break  # No need to check further if a page break is found

    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER



def remove_header_footer(doc):
    # Disable different headers and footers for odd and even pages
    doc.settings.odd_and_even_pages_header_footer = False

    # Function to clear all elements in a header or footer
    def clear_element(element):
        for child in element._element.getchildren():
            element._element.remove(child)

    # Remove header
    for section in doc.sections:
        clear_element(section.header)

    # Remove footer
    for section in doc.sections:
        clear_element(section.footer)



def start_each_heading_from_new_page(doc):
    paragraphs = list(doc.paragraphs)
    i = len(paragraphs) - 1
    while i >= 0:
        paragraph = paragraphs[i]
        if is_heading(paragraph):
            lines = paragraph.text.splitlines()
            heading_line = lines[0]

            # Insert a page break before the heading
            page_break_paragraph = paragraph.insert_paragraph_before()
            page_break_paragraph.add_run().add_break(WD_BREAK.PAGE)

            # Create a new paragraph for the heading line
            heading_paragraph = page_break_paragraph.insert_paragraph_before()
            heading_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            heading_paragraph.style = 'Heading 1'
            heading_run = heading_paragraph.add_run(heading_line.replace("\t", ""))
            heading_run.bold = True
            heading_run.font.size = Pt(14)  # Adjust font size as needed
            heading_run.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black

            # Clear the original paragraph
            paragraph.clear()

            # Insert remaining lines after the heading
            for line in reversed(lines[1:]):  # Insert in reverse to maintain order
                if line.strip():  # Skip empty lines
                    new_paragraph = paragraph.insert_paragraph_before(line.replace("\t", ""))
                    if is_subheading_heading(new_paragraph):
                        new_paragraph.style = 'Heading 2'
                        for run in new_paragraph.runs:
                            run.bold = True
                            run.font.size = Pt(12)
                        # subheading = new_paragraph.add_run(line)
                        # subheading.bold = True
                        # subheading.font.size = Pt(12)  # Adjust font size as needed
                        # subheading.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black
                    else:
                        new_paragraph.style = 'Normal'

            # Check if the next paragraph is a subheading and handle it
            if i + 1 < len(paragraphs) and is_subheading_heading(paragraphs[i + 1]):
                subheading_paragraph = paragraphs[i + 1]
                subheading_lines = subheading_paragraph.text.splitlines()
                subheading_line = subheading_lines[0]

                # Create a new paragraph for the subheading
                new_subheading_paragraph = paragraph.insert_paragraph_before()
                new_subheading_paragraph.style = 'Heading 2'
                subheading_run = new_subheading_paragraph.add_run(subheading_line)
                subheading_run.bold = True
                subheading_run.font.size = Pt(12)  # Adjust font size as needed
                subheading_run.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black

                subheading_paragraph.clear()

                # Handle the lines after the subheading
                for line in reversed(subheading_lines[1:]):  # Insert in reverse to maintain order
                    if line.strip():  # Skip empty lines
                        new_paragraph = subheading_paragraph.insert_paragraph_before(line)
                        new_paragraph.style = 'Normal'

        i -= 1



def set_page_size_to_a4(doc):
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)

    # Ensure all sections have the same page size
    for section in doc.sections:
        section.start_type = WD_SECTION_START.NEW_PAGE



def add_header_with_image_size(doc, image_path, width_cm, height_cm):
    section = doc.sections[0]  # Get the first section of the document
    doc.sections[0].header_distance  = 362585
    # Add header
    header = section.header
    header_paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()

    # Add image to header with specific width and height
    run = header_paragraph.add_run()
    picture = run.add_picture(image_path)
    picture.width = Cm(width_cm)
    picture.height = Cm(height_cm)

    # Align header content
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT




def add_footer_with_page_number(doc):
    if len(doc.sections) == 0:
        doc.add_section()

    # Set the starting page number for each section to 1
    for section in doc.sections:
        section.start_type = WD_SECTION_START.NEW_PAGE
        section.footer.is_linked_to_previous = False
        section.start_type = WD_SECTION_START.NEW_PAGE

        footer = section.footer
        footer_paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()

        # Add page number field to footer
        run = footer_paragraph.add_run()
        fld_char_begin = OxmlElement('w:fldChar')
        fld_char_begin.set(qn('w:fldCharType'), 'begin')
        run._r.append(fld_char_begin)

        instr_text = OxmlElement('w:instrText')
        instr_text.set(qn('xml:space'), 'preserve')
        instr_text.text = "PAGE \* Arabic \* MERGEFORMAT"
        run._r.append(instr_text)

        fld_char_end = OxmlElement('w:fldChar')
        fld_char_end.set(qn('w:fldCharType'), 'end')
        run._r.append(fld_char_end)

        # Align page number to the right
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # Start page numbering from 1 for each section
        sect_pr = section._sectPr
        pg_num_type = OxmlElement('w:pgNumType')
        pg_num_type.set(qn('w:start'), '1')
        sect_pr.append(pg_num_type)
        section.footer_distance = 257200



def add_page_border(doc, border_space=15):
    for section in doc.sections:
        section_start = section._sectPr
        border_element = OxmlElement('w:pgBorders')
        border_element.set(qn('w:offsetFrom'), 'page')

        for border_type in ['top', 'bottom', 'left', 'right']:
            border = OxmlElement(f'w:{border_type}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '10')
            border.set(qn('w:space'), str(border_space))  # Set the space between border and page edge
            border.set(qn('w:color'), 'auto')
            border_element.append(border)

        section_start.append(border_element)



def delete_element(element):
    """
    Delete a given element (paragraph or table) from the document.
    """
    element.getparent().remove(element)





def remove_headings_with_content(doc, headings):
    """
    Remove multiple headings and their content from the document.
    """
    delete = False
    elements_to_delete = []
    headings_to_remove = set(headings)

    for element in doc.element.body:
        # Check if the element is a paragraph
        if element.tag.endswith('p'):
            for paragraph in doc.paragraphs:
                if paragraph._element == element:
                    if is_heading(paragraph) and any(heading in paragraph.text for heading in headings_to_remove):
                        delete = True
                    elif is_heading(paragraph):
                        delete = False
                    break
        if delete:
            elements_to_delete.append(element)

    for element in elements_to_delete:
        delete_element(element)

def prompt_for_headings_to_remove(headings):
    print("Available Headings:")
    for idx, heading in enumerate(headings, start=1):
        print(f"{idx}. {heading}")

    choices = input("Enter the numbers of the headings you want to remove separated by commas (e.g., '1, 3') or enter '0' to skip: ")
    if choices.strip() == '0':
        return []

    selected_indices = []
    try:
        selected_indices = [int(idx.strip()) for idx in choices.split(',')]
        invalid_indices = [idx for idx in selected_indices if idx < 1 or idx > len(headings)]
        if invalid_indices:
            print(f"Invalid choices: {', '.join(map(str, invalid_indices))}. Please enter valid numbers.")
            return prompt_for_headings_to_remove(headings)
        return [headings[idx - 1] for idx in selected_indices]
    except ValueError:
        print("Invalid input. Please enter numbers separated by commas.")
        return prompt_for_headings_to_remove(headings)




def replace_heading_numberingssssss(doc):
    main_number = 0
    sub_number = 0
    sub_sub_number = 0
    current_main_number = 0
    current_sub_number = 0
    current_sub_sub_number = 0

    for paragraph in doc.paragraphs:
        if is_heading(paragraph):
            # Determine heading level and adjust numbering accordingly
            if re.match(r'^\d+(\.| )', paragraph.text.strip()):
                current_main_number += 1
                main_number = current_main_number  # Update main_number to current main heading number
                sub_number = 0  # Reset sub-number for new main heading
            elif re.match(r'^\d+\.\d+', paragraph.text.strip()):  # Detecting subheading format like 1.1, 1.2
                current_sub_number += 1
                sub_number = current_sub_number  # Update sub_number to current subheading number

            # Construct new heading text with updated numbering
            if re.match(r'^\d+(\.| )', paragraph.text.strip()):
                new_heading_text = "{}. {}".format(main_number, re.sub(r'^\d+(\.| )', '', paragraph.text.strip()))
            elif re.match(r'^\d+\.\d+', paragraph.text.strip()):  # Detecting subheading format like 1.1, 1.2
                new_heading_text = "{}.{}. {}".format(main_number, sub_number, re.sub(r'^\d+\.\d+\s*', '', paragraph.text.strip()))
            else:
                new_heading_text = paragraph.text.strip()

            # Replace heading text with updated numbering
            paragraph.text = new_heading_text
            paragraph.style = 'Heading 1'

            # Apply bold and size 14 to heading
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(0, 255, 0) 

        elif is_subheading(paragraph):
            # If it's a subheading (e.g., 1.1, 1.2, ...), update its numbering
            sub_number += 1  # Increment sub_number within the current main heading
            current_sub_sub_number = 0
            # Split the text to extract the heading number and the text after it
            heading_number, heading_text = re.split(r'\s+', paragraph.text.strip(), 1)

            # Split the heading text into lines
            lines = heading_text.splitlines()

            # Construct new subheading text with updated numbering for the first line only
            if lines:
                first_line_text = lines[0].strip()
                remaining_lines = "\n".join(lines[1:]) if len(lines) > 1 else ""  # Join remaining lines if any
                new_subheading_text = "{}.{}. {}".format(main_number, sub_number, first_line_text)
                paragraph.text = new_subheading_text

            # Apply style only to the first line
            paragraph.style = 'Heading 2'

            # Apply bold and size 12 to subheading for the first line only
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(255, 0, 0)  # Set font color to black

        elif is_sub_subheading(paragraph):
            current_sub_sub_number += 1
            sub_sub_number = current_sub_sub_number
            heading_number, heading_text = re.split(r'\s+', paragraph.text.strip(), 1)
            lines = heading_text.splitlines()
            if lines:
                first_line_text = lines[0].strip()
                remaining_lines = "\n".join(lines[1:]) if len(lines) > 1 else ""
                new_sub_subheading_text = "{}.{}.{}. {}".format(main_number, sub_number, sub_sub_number, first_line_text)
                paragraph.text = new_sub_subheading_text
            paragraph.style = 'Heading 3'
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0, 0, 255)

            # Append remaining lines back to the paragraph without any styling changes
            if remaining_lines:
                paragraph.add_run('\n' + remaining_lines)


def start_each_heading_from_new_line(doc):
    paragraphs = list(doc.paragraphs)
    i = len(paragraphs) - 1
    while i >= 0:
        paragraph = paragraphs[i]
        if is_subheading(paragraph):
            lines = paragraph.text.splitlines()
            heading_line = lines[0]

            # Create a new paragraph for the subheading line
            subheading_paragraph = paragraph.insert_paragraph_before(heading_line)
            subheading_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            subheading_paragraph.style = 'Heading 2'  # Assuming 'Heading 2' style for subheadings
            subheading_paragraph.runs[0].bold = True
            subheading_paragraph.runs[0].font.size = Pt(12)  # Adjust font size as needed
            subheading_paragraph.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black

            # Remove the original subheading line from the paragraph
            paragraph.clear()
            # for run in paragraph.runs:
            #     run.clear()

            # Add remaining lines as new paragraphs after the original subheading paragraph
            for line in reversed(lines[1:]):  # Insert in reverse to maintain order
                if line.strip():  # Skip empty lines
                    # Find the index of the current paragraph in the list
                    idx = paragraphs.index(paragraph)
                    # Insert new paragraph after the current paragraph
                    # new_paragraph = doc.add_paragraph(line, style='Normal')
                    p = doc.paragraphs[idx+1]
                    r = p.insert_paragraph_before('\t'+line)
        elif is_sub_subheading(paragraph):
            lines = paragraph.text.splitlines()
            heading_line = lines[0]

            # Create a new paragraph for the subheading line
            subheading_paragraph = paragraph.insert_paragraph_before(heading_line)
            subheading_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            subheading_paragraph.style = 'Heading 3'  # Assuming 'Heading 2' style for subheadings
            subheading_paragraph.runs[0].bold = True
            subheading_paragraph.runs[0].font.size = Pt(12)  # Adjust font size as needed
            subheading_paragraph.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black

            # Remove the original subheading line from the paragraph
            paragraph.clear()
            # for run in paragraph.runs:
            #     run.clear()

            # Add remaining lines as new paragraphs after the original subheading paragraph
            for line in reversed(lines[1:]):  # Insert in reverse to maintain order
                if line.strip():  # Skip empty lines
                    # Find the index of the current paragraph in the list
                    idx = paragraphs.index(paragraph)
                    # Insert new paragraph after the current paragraph
                    # new_paragraph = doc.add_paragraph(line, style='Normal')
                    p = doc.paragraphs[idx+1]
                    r = p.insert_paragraph_before('\t'+line)

        i -= 1

def remove_empty_and_excessive_spaces(doc):
    paragraphs_to_remove = []
    namespaces = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
        'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture'
    }
    for paragraph in doc.paragraphs:
        if not paragraph.text.strip():
            # Check if the paragraph has a drawing that contains a picture
            has_image = any(
                drawing.xpath('.//pic:pic')
                for drawing in paragraph._element.xpath('.//w:drawing')
            )

            # Check for tables
            has_table = paragraph._element.xpath('.//w:tbl')

            # If the paragraph is empty, does not contain an image, and does not contain a table, mark it for removal
            if not has_image and not has_table:
                paragraphs_to_remove.append(paragraph)

    # Remove identified paragraphs
    for paragraph in paragraphs_to_remove:
        p_element = paragraph._element
        p_element.getparent().remove(p_element)

    # Remove excessive spaces in the text
    for paragraph in doc.paragraphs:
        if len(paragraph.runs) > 0:
            for run in paragraph.runs:
                text = run.text
                if '   ' in text:
                    modified_text = ' '.join(text.split())
                    run.text = modified_text


def start_each_heading1_from_new_page(doc):
    first_heading_encountered = False
    paragraphs = list(doc.paragraphs)
    last_heading_index = -1

    # Identify the index of the last heading
    for i, paragraph in enumerate(paragraphs):
        if is_heading(paragraph):
            last_heading_index = i

    for i in range(len(paragraphs)):
        first_heading_encountered = False
        paragraph = paragraphs[i]
        if is_heading(paragraph):
                    # Insert a page break before the heading
                    paragraph.insert_paragraph_before().add_run().add_break(WD_BREAK.PAGE)


# Water Function

def add_watermark_to_pdf(input_pdf_path, output_pdf_path):
    watermark_text = "FLEXXON CONFIDENTIAL"

    # Create watermark
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4

    # Calculate space needed for watermark text
    watermark_font_size = 58
    text_width = c.stringWidth(watermark_text, "Helvetica", watermark_font_size)

    # Draw watermark
    c.setFillColorRGB(0.8, 0.8, 0.8, alpha=0.1)
    c.translate(width / 2, height / 2)
    c.rotate(45)
    x_position = -text_width / 2+350
    y_position = -watermark_font_size / 2

    c.setFont("Helvetica", watermark_font_size)
    c.setLineWidth(5)
    c.setStrokeColor(black)
    c.drawCentredString(x_position, y_position, watermark_text)

    c.setFillColorRGB(0, 0, 0, alpha=0.23)
    c.drawCentredString(x_position, y_position, watermark_text)

    c.showPage()
    c.save()
    buffer.seek(0)
    watermark = PdfReader(buffer)

    # Open the input PDF
    input_pdf = PdfReader(input_pdf_path)
    output_pdf = PdfWriter()

    # Add watermark to each page
    for page in input_pdf.pages:
        page.merge_page(watermark.pages[0])
        output_pdf.add_page(page)

    # Save the output PDF
    with open(output_pdf_path, "wb") as output_file:
        output_pdf.write(output_file)



def add_custom_page_at_start(input_pdf, output_pdf, logo_path,version):
    reader = PdfReader(input_pdf)

    first_page = reader.pages[0]
    width = float(first_page.mediabox.width)
    height = float(first_page.mediabox.height)

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=(width, height))

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'Title',
        parent=styles['Heading1'],
        fontSize=20,
        alignment=1,  # Center alignment
        textColor= black,
    )
    subtitle_style = ParagraphStyle(
        'Subtitle',
        parent=styles['Heading2'],
        fontSize=14,
        alignment=1,  # Center alignment
        textColor= black,
    )
    normal_style = ParagraphStyle(
        'Normal',
        parent=styles['Normal'],
        fontSize=10,
        alignment=0,  # Left alignment
        leading=14,  # Line spacing
    )

    footer_style = ParagraphStyle(
        'Footer',
        parent=styles['Normal'],
        fontSize=10,
        alignment=0,  # Center alignment
        textColor= black,
    )

    img = Image(logo_path, width=200, height=60)  # Adjust size as needed

    elements = [
        img,
        Spacer(1, 70),
        Paragraph("Industrial microSD 3.0 Specification", title_style),
        Spacer(1, 12),
        Paragraph("(FxPrem II Series, MLC)", subtitle_style),
        Spacer(1, 36),
        Paragraph("Version "+str(version), subtitle_style),
        Spacer(1, 270),
        Paragraph("Address: 28 Genting Lane, #09-03/04/05 Platinum 28, Singapore 349585", normal_style),
        Paragraph("Tel : +65-6493 5035", normal_style),
        Paragraph("Fax : +65-6493 5037", normal_style),
        Paragraph("Website: http://www.flexxon.com", normal_style),
        Paragraph("Email: flexxon@flexxon.com", normal_style),
    ]

    def add_footer(canvas, doc):
        canvas.saveState()
        footer_text = "ALL RIGHTS ARE STRICTLY RESERVED. ANY PORTION OF THIS PAPER SHALL NOT BE REPRODUCED, COPIED, OR TRANSLATED TO ANY OTHER FORMS WITHOUT PERMISSION FROM FLEXXON."
        footer = Paragraph(footer_text, footer_style)
        w, h = footer.wrap(doc.width, doc.bottomMargin)
        footer.drawOn(canvas, doc.leftMargin, h)
        canvas.restoreState()

    # Build the PDF with the footer
    doc.build(elements, onFirstPage=add_footer, onLaterPages=add_footer)

    # Create a PDF writer object
    writer = PdfWriter()

    # Add the custom page
    custom_pdf = PdfReader(BytesIO(buffer.getvalue()))
    writer.add_page(custom_pdf.pages[0])

    # Add all pages from the existing PDF
    for page in reader.pages:
        writer.add_page(page)

    # Write the output to a file
    with open(output_pdf, "wb") as output_file:
        writer.write(output_file)



def remove_header_footer_pdf(input_pdf, output_pdf, header_height, footer_height):
    try:
        pdf_document = fitz.open(input_pdf)

        # Iterate through each page
        for page_num in range(len(pdf_document)):
            page = pdf_document.load_page(page_num)
            rect = page.rect

            # Define new crop box to remove header and footer
            new_rect = fitz.Rect(
                rect.x0,
                rect.y0 + header_height,
                rect.x1,
                rect.y1 - footer_height
            )

            # Set the new crop box
            page.set_cropbox(new_rect)

        # Save the modified PDF
        pdf_document.save(output_pdf)
        print(f"Successfully processed {input_pdf} and saved as {output_pdf}")
        return output_pdf
    except Exception as e:
        print(f"An error occurred while processing the PDF: {e}")
        return f"An error occurred while processing the PDF: {e}"


def create_index_of_heading(document):
    # Create a new paragraph for the heading "Table of Contents"
    toc_heading = document.add_paragraph()
    toc_run = toc_heading.add_run("Table of Contents")
    toc_run.bold = True
    toc_run.font.size = docx.shared.Pt(16)  # Adjust font size if needed
    toc_run.font.name = 'Arial'  # Adjust font family if needed
    toc_heading.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    toc_heading.space_after = docx.shared.Pt(12)  # Adjust spacing after heading

    # Create a new paragraph for the TOC in a temporary document
    temp_doc = docx.Document()
    toc_paragraph = temp_doc.add_paragraph()
    run = toc_paragraph.add_run()

    # Define the TOC field elements
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:updateFields')
    fldChar3.set(qn('w:val'), 'true')

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    # Append the TOC field elements to the run element
    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar3)
    r_element.append(fldChar4)

    # Get the XML of the new paragraph
    toc_xml = toc_paragraph._element

    # Extract the body element of the main document
    body = document._element.body

    # Insert the heading and TOC paragraph XML at the beginning of the body
    body.insert(0, toc_heading._element)
    body.insert(1, toc_xml)



def convert_docx_to_pdf(docx_path, pdf_path):
    try:
        # Create an instance of Word application
        word_app = win32com.client.Dispatch("Word.Application")
        word_app.Visible = False # Run in the background        
        # Open the DOCX file 
        doc = word_app.Documents.Open(os.path.abspath(docx_path))# Save as PDF        
        doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17) # 17 is the wdFormatPDF constant        
        doc.Close() # Quit the Word application        
        word_app.Quit()
        print(f"Conversion successful: {docx_path} to {pdf_path}") 
    except Exception as e:
        print(f"An error occurred: {e}")# Usage 







def convert_docx_to_pdf_windows(pdf_path,Foldername):
    word = win32com.client.Dispatch("Word.Application")
    word.visible = 0  # Change to 1 if you want to see Word application running

    # Get file name and normalized path
    filename = os.path.basename(pdf_path)
    in_file = os.path.abspath(pdf_path)

    # Fixed output directory for DOCX files
    output_dir = r"F:\\pdf_extract\\"+Foldername  # Replace with your desired output directory

    # Generate full path for DOCX file
    docx_filename = filename[:-4] + ".docx"  # Remove .pdf extension and add .docx
    docx_path = os.path.join(output_dir, docx_filename)

    # Convert PDF to DOCX
    wb = word.Documents.Open(in_file)
    wb.SaveAs2(docx_path, FileFormat=16)
    wb.Close()
    word.Quit()

    return docx_path

def add_spacing_after_paragraph(paragraph,space = '200'):
    p_pr = paragraph._element.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:after'), space)  # Space after paragraph in twips (200 twips = 0.25 inches)

    p_pr.append(spacing)

def add_bottom_border(paragraph):
    p_pr = paragraph._element.get_or_add_pPr()
    borders = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')  # Border type (e.g., 'single', 'double', 'dotted', etc.)
    bottom.set(qn('w:sz'), '15')       # Border size (in eighths of a point, so '4' means 0.5pt)
    bottom.set(qn('w:space'), '8')    # Space between border and text
    bottom.set(qn('w:color'), '000000') # Border color in hex (black in this case)

    borders.append(bottom)
    p_pr.append(borders)
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:after'), '200')  # Space after paragraph in twips (200 twips = 0.25 inches)
    p_pr.append(spacing)


def replace_heading_numbering(doc):
    main_number = 0
    sub_number = 0
    sub_sub_number = 0
    current_main_number = 0
    current_sub_number = 0
    current_sub_sub_number = 0

    for paragraph in doc.paragraphs:
        if is_heading(paragraph):
            current_main_number += 1
            main_number = current_main_number
            current_sub_number = 0
            current_sub_sub_number = 0
            new_heading_text = "{}. {}".format(main_number, re.sub(r'^\d+(\.| )', '', paragraph.text.strip()))
            paragraph.text = new_heading_text
            paragraph.style = 'Heading 1'
            add_bottom_border(paragraph)
            for run in paragraph.runs:
                # run.bold = True
                run.font.size = Pt(16)
                run.font.color.rgb = RGBColor(0, 0, 0)

        elif is_subheading(paragraph):
            current_sub_number += 1
            sub_number = current_sub_number
            current_sub_sub_number = 0
            heading_number, heading_text = re.split(r'\s+', paragraph.text.strip(), 1)
            lines = heading_text.splitlines()
            if lines:
                first_line_text = lines[0].strip()
                remaining_lines = "\n".join(lines[1:]) if len(lines) > 1 else ""
                new_subheading_text = "\t{}.{}. {}".format(main_number, sub_number, first_line_text)
                paragraph.text = new_subheading_text
            paragraph.style = 'Heading 2'
            add_spacing_after_paragraph(paragraph)
            for run in paragraph.runs:
                # run.bold = True
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(0, 0, 0)
            if remaining_lines:
                paragraph.add_run('\n' + remaining_lines)

        elif is_sub_subheading(paragraph):
            current_sub_sub_number += 1
            sub_sub_number = current_sub_sub_number
            heading_number, heading_text = re.split(r'\s+', paragraph.text.strip(), 1)
            lines = heading_text.splitlines()
            if lines:
                first_line_text = lines[0].strip()
                remaining_lines = "\n".join(lines[1:]) if len(lines) > 1 else ""
                new_sub_subheading_text = "\t\t{}.{}.{}. {}".format(main_number, sub_number, sub_sub_number, first_line_text)
                paragraph.text = new_sub_subheading_text
            paragraph.style = 'Heading 3'
            add_spacing_after_paragraph(paragraph)
            for run in paragraph.runs:
                # run.bold = True
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0, 0, 0)
            if remaining_lines:
                paragraph.add_run('\n' + remaining_lines)


def remove_line_under_heading1(doc):
    for paragraph in doc.paragraphs:
        if paragraph.style.name == 'Heading 1':
            for run in paragraph.runs:
                for shape in run.inline_shapes:
                    if shape.type == 3:  # Check if it's a line shape (type 3)
                        # Remove the line shape
                        shape._element.getparent().remove(shape._element)


def update_toc_with_win32(doc_path):
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False

    doc = word.Documents.Open(os.path.abspath(doc_path))
    doc.TablesOfContents(1).Update()
    doc.Save()
    doc.Close()
    word.Quit()
